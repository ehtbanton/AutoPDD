import docx
import os
import shutil
import json
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml import OxmlElement # <--- ADD THIS IMPORT
from docx.text.paragraph import Paragraph # <--- ADD THIS IMPORT
import sys

# --- HELPER FUNCTIONS ---
def _iter_block_items(parent):
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, docx.table._Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Parent must be a Document or _Cell object")
    
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield docx.table.Table(child, parent)

def load_word_doc_to_string(folder_path):
    filename = None
    if not os.path.exists(folder_path) or not os.path.isdir(folder_path):
         return f"Error: Directory not found or is not a directory at '{folder_path}'"
    try:
        for f in os.listdir(folder_path):
            if f.lower().endswith('.docx') and not f.startswith('~$'):
                filename = os.path.join(folder_path, f)
                break
    except FileNotFoundError:
        return f"Error: Directory not found at '{folder_path}'"
    if not filename:
        return f"Error: No .docx file found in the directory '{folder_path}'"
    try:
        document = docx.Document(filename)
        full_text_blocks = []
        for block in _iter_block_items(document):
            if isinstance(block, docx.text.paragraph.Paragraph):
                if block.text.strip():
                    full_text_blocks.append(block.text)
            elif isinstance(block, docx.table.Table):
                if not block.rows: continue
                table_lines = ["| " + " | ".join(cell.text.replace('\n', ' ').strip() for cell in block.rows[0].cells) + " |"]
                table_lines.append("| " + " | ".join(['---'] * len(block.rows[0].cells)) + " |")
                for row in block.rows[1:]:
                    table_lines.append("| " + " | ".join(cell.text.replace('\n', ' ').strip() for cell in row.cells) + " |")
                full_text_blocks.append("\n".join(table_lines))
        return "\n\n".join(full_text_blocks)
    except Exception as e:
        return f"Error processing file '{os.path.basename(filename)}': {e}"

def create_output_doc_from_template(project_name):
    script_dir = os.path.dirname(os.path.abspath(__file__))
    backend_dir = os.path.abspath(os.path.join(script_dir, '..'))
    template_folder = os.path.join(backend_dir, "pdd_template")
    output_folder = os.path.join(backend_dir, "auto_pdd_output")
    
    if not os.path.isdir(template_folder):
        os.makedirs(template_folder)
        print(f"Created template directory: {template_folder}")
        raise FileNotFoundError(f"Error: Template directory was missing. Please upload a template .docx file to this folder and try again.")

    template_path = next((os.path.join(template_folder, f) for f in os.listdir(template_folder) if f.lower().endswith('.docx') and not f.startswith('~$')), None)
    
    if not template_path:
        raise FileNotFoundError(f"Error: No .docx template found in '{template_folder}'")

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    output_path = os.path.join(output_folder, f"AutoPDD_{project_name}.docx")
    if not os.path.exists(output_path):
        shutil.copy(template_path, output_path)
        print(f"Created output document at: {output_path}")
    else:
        print(f"Output document already exists at: {output_path}. This file will be updated.")
    return output_path

def replace_section_in_word_doc(doc_path, start_marker, end_marker, ai_markdown_text, status):
    doc = docx.Document(doc_path)
    all_blocks = list(_iter_block_items(doc))

    start_index, end_index = -1, len(all_blocks)
    for i, block in enumerate(all_blocks):
        if isinstance(block, docx.text.paragraph.Paragraph):
            if block.text.strip() == start_marker:
                start_index = i
            elif start_index != -1 and block.text.strip() == end_marker:
                end_index = i
                break
    
    if start_index == -1:
        print(f"  > WARNING: Start marker '{start_marker}' not found.")
        return

    # Delete existing content
    for i in range(end_index - 1, start_index, -1):
        element = doc.element.body[i]
        doc.element.body.remove(element)

    # This will require a robust markdown parsing logic
    # and then creating new docx elements.
    # ...

    doc.save(doc_path)


def replace_paragraph_text_by_index(doc, paragraph_index, new_text):
    """
    Finds and replaces the text in a paragraph by its index, preserving the original formatting.
    This function iterates through all paragraphs in the document, including those within tables,
    counting them in order to find the paragraph at the specified index.
    """
    current_index = 0

    # Iterate through all blocks in the document using the existing helper function
    for block in _iter_block_items(doc):
        if isinstance(block, docx.text.paragraph.Paragraph):
            if current_index == paragraph_index:
                # Found the target paragraph - clear its existing content
                for run in block.runs:
                    run.clear()
                # Add a new run with the new text, preserving the paragraph's style
                block.add_run(new_text)
                return True # Indicate success
            current_index += 1
        elif isinstance(block, docx.table.Table):
            # Iterate through paragraphs in table cells
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if current_index == paragraph_index:
                            # Found the target paragraph - clear its existing content
                            for run in p.runs:
                                run.clear()
                            # Add a new run with the new text, preserving the paragraph's style
                            p.add_run(new_text)
                            return True # Indicate success
                        current_index += 1

    return False # Indicate that the paragraph index was not found


def replace_paragraph_text(doc, old_text, new_text):
    """
    Finds and replaces the text in a paragraph, preserving the original formatting.
    This function iterates through all paragraphs in the document, including those within tables.
    (DEPRECATED: Use replace_paragraph_text_by_index for more reliable paragraph identification)
    """
    for p in doc.paragraphs:
        if p.text.strip() == old_text.strip():
            # Clear existing runs in the paragraph
            for run in p.runs:
                run.clear()
            # Add a new run with the new text, preserving the paragraph's style
            p.add_run(new_text)
            return True # Indicate success

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip() == old_text.strip():
                        for run in p.runs:
                            run.clear()
                        p.add_run(new_text)
                        return True # Indicate success
    return False # Indicate that the text was not found

if __name__ == '__main__':
    # For debugging: print the received arguments
    # print(f"Received arguments: {sys.argv}")

    if len(sys.argv) != 4:
        print("Usage: python word_editor.py <doc_path> <paragraph_index> <new_text>")
        print(f"Error: Expected 4 arguments, but received {len(sys.argv)}.")
        sys.exit(1)

    doc_path = sys.argv[1]
    try:
        paragraph_index = int(sys.argv[2])
    except ValueError:
        print(f"Error: paragraph_index must be a valid integer, received: '{sys.argv[2]}'")
        sys.exit(1)
    new_text = sys.argv[3]

    if not os.path.exists(doc_path):
        print(f"Error: Document not found at '{doc_path}'")
        sys.exit(1)

    try:
        document = docx.Document(doc_path)
        if replace_paragraph_text_by_index(document, paragraph_index, new_text):
            document.save(doc_path)
            print("SUCCESS")
        else:
            print(f"Error: Could not find paragraph at index: {paragraph_index}")
            sys.exit(1)
    except Exception as e:
        print(f"An error occurred: {e}")
        sys.exit(1)