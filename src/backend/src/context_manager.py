
import os
import json
import pdfplumber
import docx
import sys

def _extract_text_from_file(file_path):
    """
    A helper function to extract text and tables from a single file.

    Args:
        file_path (str): The full path to the .pdf or .docx file.

    Returns:
        tuple: (content_string, file_metadata)
            content_string: The extracted text content, with tables in Markdown format.
            file_metadata: Dict with file path and page count info for linking
    """
    content_parts = []
    filename = os.path.basename(file_path)
    file_metadata = {
        'full_path': file_path,
        'filename': filename,
        'total_pages': 0
    }

    try:
        # --- Handle PDF files ---
        if filename.lower().endswith('.pdf'):
            with pdfplumber.open(file_path) as pdf:
                file_metadata['total_pages'] = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    page_number = i + 1
                    page_text = page.extract_text()
                    if page_text:
                        # Add page marker for reference tracking
                        content_parts.append(f"\n--- PAGE {page_number} START ---\n{page_text}\n--- PAGE {page_number} END ---\n")

                    # Extract tables and convert to Markdown
                    tables = page.extract_tables()
                    for table in tables:
                        if not table: continue
                        header = "| " + " | ".join(str(cell) if cell is not None else '' for cell in table[0]) + " |"
                        separator = "| " + " | ".join(["---"] * len(table[0])) + " |"
                        rows = ["| " + " | ".join(str(cell) if cell is not None else '' for cell in row) + " |" for row in table[1:]]
                        markdown_table = "\n".join([header, separator] + rows)
                        content_parts.append(f"\n\n--- Table on Page {page_number} ---\n{markdown_table}\n")

        # --- Handle Word (.docx) files ---
        elif filename.lower().endswith('.docx'):
            doc = docx.Document(file_path)
            file_metadata['total_pages'] = 1  # Word docs don't have clear page boundaries

            # Estimate page numbers based on content length (rough approximation)
            paragraph_count = 0
            for para in doc.paragraphs:
                paragraph_count += 1
                if para.text.strip():
                    # Rough page estimation: ~25 paragraphs per page
                    estimated_page = (paragraph_count // 25) + 1
                    content_parts.append(f"[Page {estimated_page}] {para.text}")

            # Extract tables and convert to Markdown
            for i, table in enumerate(doc.tables):
                if not table.rows: continue
                header_cells = table.rows[0].cells
                header = "| " + " | ".join(cell.text.strip() for cell in header_cells) + " |"
                separator = "| " + " | ".join(["---"] * len(header_cells)) + " |"
                rows = ["| " + " | ".join(cell.text.strip() for cell in row.cells) + " |" for row in table.rows[1:]]
                markdown_table = "\n".join([header, separator] + rows)
                content_parts.append(f"\n\n--- Table {i+1} ---\n{markdown_table}\n")

    except Exception as e:
        print(f"Could not process file '{filename}'. Reason: {e}")
        sys.stdout.flush()
        return "", file_metadata # Return empty string and metadata on failure

    return "\n".join(content_parts), file_metadata


def extract_text_from_folder(folder_path):
    """
    Extracts text from PDF and Word files in a folder and maintains a TXT
    file containing the content in a structured (JSON) format, updating it 
    with any new or deleted files.

    Args:
        folder_path (str): The absolute or relative path to the folder.

    Returns:
        bool: True if the TXT file was modified (files added/removed),
              False otherwise.
    """
    if not os.path.isdir(folder_path):
        print(f"Error: Folder not found at '{folder_path}'")
        sys.stdout.flush()
        # Let's create it
        os.makedirs(folder_path)
        print(f"Created directory: {folder_path}")
        sys.stdout.flush()


    # The only change needed is the file extension
    txt_filepath = os.path.join(folder_path, "all_context.txt")
    changes_made = False

    # 1. Load existing data from all_context.txt or create an empty list
    try:
        with open(txt_filepath, 'r', encoding='utf-8') as f:
            # Read the file's text content, then parse it as JSON
            all_context = json.loads(f.read())
    except (FileNotFoundError, json.JSONDecodeError):
        all_context = []
    
    # Get a set of filenames we already have processed
    known_files = {entry['filename'] for entry in all_context}

    # Get a set of current .pdf and .docx files in the folder
    current_files = {
        f for f in os.listdir(folder_path) 
        if f.lower().endswith(('.pdf', '.docx'))
    }

    # 2. Handle deletions
    files_to_remove = known_files - current_files
    if files_to_remove:
        print(f"Files removed: {', '.join(files_to_remove)}")
        sys.stdout.flush()
        all_context = [
            entry for entry in all_context 
            if entry['filename'] not in files_to_remove
        ]
        changes_made = True

    # 3. Handle additions
    files_to_add = current_files - known_files
    if files_to_add:
        print(f"New files found: {', '.join(files_to_add)}")
        sys.stdout.flush()
        for filename in files_to_add:
            file_path = os.path.join(folder_path, filename)
            print(f"-> Processing: {filename}")
            sys.stdout.flush()
            
            text_content, file_metadata = _extract_text_from_file(file_path)

            if text_content:
                all_context.append({
                    'filename': filename,
                    'text_content': text_content,
                    'file_metadata': file_metadata
                })
                print(f"   ...extracted {len(text_content)} characters.")
                sys.stdout.flush()
                changes_made = True

    # 4. Save the updated data back to the TXT file if any changes were made
    if changes_made:
        print(f"\nSaving changes to '{txt_filepath}'...")
        sys.stdout.flush()
        try:
            with open(txt_filepath, 'w', encoding='utf-8') as f:
                # Convert the Python list to a JSON-formatted string
                json_string = json.dumps(all_context, indent=4)
                # Write that string to the .txt file
                f.write(json_string)
            print("...Success!")
            sys.stdout.flush()
        except Exception as e:
            print(f"Error saving the TXT file: {e}")
            sys.stdout.flush()
            return False
    else:
        print("\nNo changes detected. Content is up-to-date.")
        sys.stdout.flush()

    return changes_made
