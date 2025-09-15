#!/usr/bin/env python3

import sys
import json
import docx
from typing import Dict, Any, Optional

def update_by_offset(doc_path: str, offset_json: str, new_text: str) -> bool:
    """
    Update document content using structural offset instead of text matching.

    Args:
        doc_path: Path to the Word document
        offset_json: JSON string with offset information
        new_text: New text to insert

    Returns:
        True if successful, False otherwise
    """
    try:
        # Parse offset information
        offset = json.loads(offset_json)

        # Load document
        doc = docx.Document(doc_path)

        print(f"Updating document using offset: {offset}")
        print(f"New text: '{new_text}'")

        # Handle table cell updates
        if 'tableIndex' in offset and offset['tableIndex'] is not None:
            return update_table_cell(doc, offset, new_text)

        # Handle paragraph updates
        elif 'paragraphIndex' in offset and offset['paragraphIndex'] >= 0:
            return update_paragraph_by_index(doc, offset, new_text)

        else:
            print("Error: Invalid offset structure")
            return False

    except Exception as e:
        print(f"Error updating document: {e}")
        return False

def update_table_cell(doc, offset: Dict[str, Any], new_text: str) -> bool:
    """Update a specific table cell using offset coordinates"""
    try:
        table_index = offset['tableIndex']
        row_index = offset['rowIndex']
        cell_index = offset['cellIndex']

        print(f"Targeting table {table_index}, row {row_index}, cell {cell_index}")

        # Get all tables
        tables = doc.tables

        if table_index >= len(tables):
            print(f"Error: Table index {table_index} out of range (only {len(tables)} tables)")
            return False

        table = tables[table_index]

        if row_index >= len(table.rows):
            print(f"Error: Row index {row_index} out of range (only {len(table.rows)} rows)")
            return False

        row = table.rows[row_index]

        if cell_index >= len(row.cells):
            print(f"Error: Cell index {cell_index} out of range (only {len(row.cells)} cells)")
            return False

        cell = row.cells[cell_index]

        # Clear existing content and add new content
        # Get the first paragraph in the cell
        if cell.paragraphs:
            paragraph = cell.paragraphs[0]

            # Clear existing runs
            for run in paragraph.runs:
                run.clear()

            # Add new content
            paragraph.add_run(new_text)

            print(f"Successfully updated table cell at [{table_index}][{row_index}][{cell_index}]")

            # Save document
            doc.save(doc_path)
            return True
        else:
            print("Error: No paragraphs found in target cell")
            return False

    except Exception as e:
        print(f"Error updating table cell: {e}")
        return False

def update_paragraph_by_index(doc, offset: Dict[str, Any], new_text: str) -> bool:
    """Update a specific paragraph using its index"""
    try:
        paragraph_index = offset['paragraphIndex']

        print(f"Targeting paragraph {paragraph_index}")

        # Get all paragraphs
        paragraphs = doc.paragraphs

        if paragraph_index >= len(paragraphs):
            print(f"Error: Paragraph index {paragraph_index} out of range (only {len(paragraphs)} paragraphs)")
            return False

        paragraph = paragraphs[paragraph_index]

        # Clear existing runs
        for run in paragraph.runs:
            run.clear()

        # Add new content
        paragraph.add_run(new_text)

        print(f"Successfully updated paragraph at index {paragraph_index}")

        # Save document
        doc.save(doc_path)
        return True

    except Exception as e:
        print(f"Error updating paragraph: {e}")
        return False

def main():
    """Main function for command line usage"""
    if len(sys.argv) != 4:
        print("Usage: python word_editor_offset.py <doc_path> <offset_json> <new_text>")
        sys.exit(1)

    doc_path = sys.argv[1]
    offset_json = sys.argv[2]
    new_text = sys.argv[3]

    try:
        success = update_by_offset(doc_path, offset_json, new_text)

        if success:
            print("SUCCESS")
            sys.exit(0)
        else:
            print("FAILURE")
            sys.exit(1)

    except Exception as e:
        print(f"FAILURE: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()